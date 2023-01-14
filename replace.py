# 导入模块

import os
import sys

from docx import Document


def replace_word(doc, old_word, new_word):
    
    #param doc: 要替换的文档
    #param old_word: 被替换的文字
    #param new_word: 替换后的文字
    
    for p in doc.paragraphs:  # 遍历文档段落
        for run in p.runs:  # 遍历段落的字块
            run.text = run.text.replace(old_word, new_word)  # 替换字块的文字，然后赋值给字块

    # 遍历文档的表格， 替换表格里的要替换的文字
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_word, new_word)



path = sys.argv[1]
old_word = sys.argv[2]
new_word = sys.argv[3]

for root, dirs, files in os.walk(path):
    for file in files:

        file_name = os.path.join(root, file)
        ext = os.path.splitext(file_name)[-1]
        
        if (ext == '.docx'):

            print(file_name)

            doc = Document(file_name)

            # 执行替换函数
            replace_word(doc, old_word, new_word)

            doc.save(file_name)






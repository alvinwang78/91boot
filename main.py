import sys 
import os
import string

from PySide6 import QtWidgets
from PySide6.QtWidgets import QApplication, QMessageBox, QMainWindow
from PySide6.QtCore import QFile

from docx import Document as Word
from openpyxl import load_workbook as Excel

from ui.main import Ui_mainWindow

# 遍历文件夹
def scan_path(path_name,file_list):

  for item in os.scandir(path_name):

      if item.is_dir():

        scan_path(item.path,file_list)

      elif item.is_file():

        ext = os.path.splitext(item.path)[-1]

        if (item.name.startswith('.') == False):
          file_list.append(item.path)

#替换word里的文字
def replace_word(file, old_word, new_word):

  doc = Word(file)

  # 执行替换函数
  for p in doc.paragraphs:  # 遍历文档段落
    for run in p.runs:  # 遍历段落的字块
      if old_word in run.text:
        run.text = run.text.replace(old_word, new_word)

  # 遍历文档的表格， 替换表格里的要替换的文字
  for table in doc.tables:
      for row in table.rows:
          for cell in row.cells:
            for p in cell.paragraphs:
              for run in p.runs:
                if old_word in run.text:
                 run.text = run.text.replace(old_word, new_word)

  # 替换页眉
  for p in doc.sections[0].header.paragraphs:
    for run in p.runs:
      if old_word in run.text:
        run.text = run.text.replace(old_word, new_word)

  # 替换页脚
  for p in doc.sections[0].footer.paragraphs:
    for run in p.runs:
      if old_word in run.text:
        run.text = run.text.replace(old_word, new_word)

  doc.save(file)

#替换excel里的文字
def replace_excel(file, old_word, new_word):

  xls = Excel(file)
  sheets = xls.get_sheet_names()

  for i in range(len(sheets)):

    sheet = xls.get_sheet_by_name(sheets[i])

    for col in sheet.iter_cols():
      for cell in col:
        if isinstance(cell.value, str):
          cell.value = cell.value.replace(old_word, new_word)

  xls.save(file)

#替换文件名
def rename_file(old_file_name, old_word, new_word):

  new_file_name = old_file_name.replace(old_word, new_word)
  os.rename(old_file_name,new_file_name)

  return new_file_name

# 主界面显示
class MainWindow (QMainWindow):
  
  def __init__(self):

    super(MainWindow, self).__init__()
 
    self.ui = Ui_mainWindow()
    self.ui.setupUi(self)
    self._banding()
 
  def _banding(self):
    self.ui.choosePath.clicked.connect(self.choosePath)
    self.ui.list.clicked.connect(self.list)
    self.ui.start.clicked.connect(self.start)

  # 选择目录
  def choosePath(self):
    path_name = QtWidgets.QFileDialog.getExistingDirectory(self, "选择文件夹","../")
    self.ui.pathName.setText(path_name)

  # 列出文件
  def list(self):
    path_name = self.ui.pathName.text()

    if path_name == '':

      log_html = '请选择目录'
      self.ui.showLog.setHtml(log_html)

    else:

      file_list = []
      scan_path(path_name,file_list)

      log_html = '开始搜索 <br><br>'
      self.ui.showLog.setHtml(log_html)

      for file in file_list:

        log_html = log_html + file.replace(path_name,'') + '<br>'
        self.ui.showLog.setHtml(log_html)

      log_html = log_html + '<br>搜索完成' + '<br>共找到 ' + str(len(file_list)) + ' 个文件'
      self.ui.showLog.setHtml(log_html)
  
  # 开始替换
  def start(self):

    self.ui.start.setText('执行中....')
    self.ui.start.setEnabled(False)

    path_name = self.ui.pathName.text()

    old_word = self.ui.oldWord.text()
    new_word = self.ui.newWord.text()

    is_rename = self.ui.isRename.isChecked()

    if path_name == '':

      log_html = '请选择目录'
      self.ui.showLog.setHtml(log_html)

    else:

      result = QMessageBox.question(self, "警告", "确定要开始执行?", QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes)

      if result == QMessageBox.Yes:

        file_list = []
        scan_path(path_name,file_list)

        log_html = '开始执行 <br><br>'
        self.ui.showLog.setHtml(log_html)

        for file in file_list:

          if os.path.splitext(file)[-1] == '.docx':
           replace_word(file,old_word,new_word)
          elif os.path.splitext(file)[-1] == '.xlsx':
           replace_excel(file,old_word,new_word)

          #替换文件名
          if is_rename is True:
            file = rename_file(file,old_word,new_word)

          log_html = log_html + file.replace(path_name,'') + '<br>'
          self.ui.showLog.setHtml(log_html)

        log_html = log_html + '<br>执行完成<br>共替换 ' + str(len(file_list)) + ' 个文件'
        self.ui.showLog.setHtml(log_html)

    self.ui.start.setText('开始替换')
    self.ui.start.setEnabled(True)

# 执行
if __name__ == "__main__":
  app = QtWidgets.QApplication(sys.argv)
 
  window = MainWindow()
  window.show()

  sys.exit(app.exec())


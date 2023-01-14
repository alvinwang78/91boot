import sys 
import os

from PySide6 import QtWidgets
from PySide6.QtWidgets import QApplication, QMainWindow
from PySide6.QtCore import QFile
 
from ui.CreateFile import Ui_createFile

from docx import Document

class MainWindow (QMainWindow):
  
  def __init__(self):
    super(MainWindow, self).__init__()
 
    self.ui = Ui_createFile()
    self.ui.setupUi(self)
    self._banding()
 
  def _banding(self):
    self.ui.chooseTemplateFolder.clicked.connect(self.chooseTemplateFolder)
    self.ui.chooseTargetFolder.clicked.connect(self.chooseTargetFolder)
    self.ui.create.clicked.connect(self.create)
    
  def chooseTemplateFolder(self):
    template_folder = QtWidgets.QFileDialog.getExistingDirectory(self, "选择文件夹","../")
    
    self.ui.templateFolder.setText(template_folder)

  def chooseTargetFolder(self):
    target_folder = QtWidgets.QFileDialog.getExistingDirectory(self, "选择文件夹","../")
    
    self.ui.targetFolder.setText(target_folder)
    
    
  def create(self):

    self.ui.create.setText('生成中....')
    self.ui.create.setEnabled(False)

    template_folder = self.ui.templateFolder.text()
    target_folder = self.ui.targetFolder.text()

    old_word = self.ui.oldWord.text()
    new_word = self.ui.newWord.text()

    if (template_folder == '' or target_folder == '' or old_word == '' or new_word == ''):

      log_html = '请选择或输入全部信息'
      self.ui.showLog.setHtml(log_html)

    else:

      log_html = '开始生成，请稍候 <br>'
      self.ui.showLog.setHtml(log_html)

      for root, dirs, files in os.walk(template_folder):
        for file in files:

          template_file = os.path.join(root, file)
          ext = os.path.splitext(template_file)[-1]

          # 创建目录
          target_root = root.replace(template_folder,target_folder)
          if os.path.exists(target_root) == False:
            os.mkdir(target_root)
          
          if (ext == '.docx'):

            target_file = template_file.replace(template_folder,target_folder)

            log_html = log_html + '文件 ' + template_file + ' > '+ target_file + '<br>'
            self.ui.showLog.setHtml(log_html)

            doc = Document(template_file)

            # 执行替换函数
            for paragraph in doc.paragraphs:  # 遍历文档段落
              for run in paragraph.runs:  # 遍历段落的字块
                print(run.text)
                run.text = run.text.replace(old_word, new_word)  # 替换字块的文字，然后赋值给字块

            # 遍历文档的表格， 替换表格里的要替换的文字
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell.text = cell.text.replace(old_word, new_word)

            doc.save(target_file)

      log_html = log_html + '恭喜您，文档已全部生成'
      self.ui.showLog.setHtml(log_html)

    self.ui.create.setText('开始生成')
    self.ui.create.setEnabled(True)

if __name__ == "__main__":
  app = QtWidgets.QApplication(sys.argv)
 
  window = MainWindow()
  window.show()

  sys.exit(app.exec())